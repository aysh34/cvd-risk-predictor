"""
Streamlit Frontend for CVD Clinical Decision Support System
Calls Elasticsearch Agent Builder API for agentic reasoning
"""

import streamlit as st
import requests
import json
import plotly.graph_objects as go
from datetime import datetime

# Configuration
ELASTIC_API_KEY = st.secrets['ELASTIC_API_KEY']

# Page config
st.set_page_config(
    page_title="CVD Risk Assessment",
    page_icon="🫀",
    layout="wide"
)
from docx import Document
from datetime import datetime
import io

if "risk_report_data" not in st.session_state:
    st.session_state.risk_report_data = None
# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .risk-high {
        background-color: #ffebee;
        color: black;
        padding: 1rem;
        border-left: 5px solid #f44336;
        border-radius: 5px;
    }
    .risk-moderate {
        background-color: #fff3e0;
        padding: 1rem;
        color: black;
        border-left: 5px solid #ff9800;
        border-radius: 5px;
    }
    .risk-low {
        background-color: #e8f5e9;
        padding: 1rem;
        color: black;
        border-left: 5px solid #4caf50;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown('<div class="main-header">🫀 Cardiovascular Risk Assessment</div>', unsafe_allow_html=True)
st.markdown("AI-powered clinical decision support with evidence-based recommendations")

# Sidebar for patient input
st.sidebar.header("📋 Patient Information")

with st.sidebar:
    st.subheader("Demographics")
    
    gender = st.selectbox(
        "Gender",
        options=[1, 2],
        format_func=lambda x: "Female" if x == 1 else "Male",
        help="Biological sex"
    )
    
    age_years = st.number_input(
        "Age (years)",
        min_value=18,
        max_value=100,
        value=55,
        step=1
    )
    
    col1, col2 = st.columns(2)
    with col1:
        height = st.number_input(
            "Height (cm)",
            min_value=140,
            max_value=220,
            value=170,
            step=1
        )
    with col2:
        weight = st.number_input(
            "Weight (kg)",
            min_value=40.0,
            max_value=200.0,
            value=80.0,
            step=0.5
        )
    
    # Auto-calculate BMI
    height_m = height / 100
    bmi = weight / (height_m ** 2)
    st.metric("Calculated BMI", f"{bmi:.1f}")
    
    st.divider()
    st.subheader("Blood Pressure")
    
    col1, col2 = st.columns(2)
    with col1:
        ap_hi = st.number_input(
            "Systolic (mmHg)",
            min_value=80,
            max_value=250,
            value=140,
            step=5,
            help="Upper number"
        )
    with col2:
        ap_lo = st.number_input(
            "Diastolic (mmHg)",
            min_value=50,
            max_value=150,
            value=90,
            step=5,
            help="Lower number"
        )
    
    # Auto-calculate derived values
    pulse_pressure = ap_hi - ap_lo
    map_value = (ap_hi + 2 * ap_lo) / 3
    
    st.caption(f"Pulse Pressure: {pulse_pressure} mmHg")
    st.caption(f"MAP: {map_value:.1f} mmHg")
    
    st.divider()
    st.subheader("Lab Values")
    
    cholesterol = st.radio(
        "Cholesterol",
        options=[1, 2, 3],
        format_func=lambda x: ["Normal", "Above Normal", "Well Above Normal"][x-1],
        horizontal=True
    )
    
    gluc = st.radio(
        "Glucose",
        options=[1, 2, 3],
        format_func=lambda x: ["Normal", "Above Normal", "Well Above Normal"][x-1],
        horizontal=True
    )
    
    st.divider()
    st.subheader("Lifestyle Factors")
    
    smoke = st.checkbox("Current Smoker", value=False)
    alco = st.checkbox("Consumes Alcohol", value=False)
    active = st.checkbox("Physically Active", value=True)


# Main area
if st.sidebar.button("🔍 Assess Risk", type="primary", width='stretch'):
    
    # Prepare patient data
    patient_data = {
        "gender": gender,
        "age_years": age_years,
        "height": height,
        "weight": weight,
        "bmi": round(bmi, 1),
        "ap_hi": ap_hi,
        "ap_lo": ap_lo,
        "pulse_pressure": pulse_pressure,
        "map": round(map_value, 1),
        "cholesterol": cholesterol,
        "gluc": gluc,
        "smoke": 1 if smoke else 0,
        "alco": 1 if alco else 0,
        "active": 1 if active else 0
    }
    
    # Show progress
    with st.spinner("🤖 Agent is working..."):
        
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Step 1: Call ML model
        status_text.text("📊 Calculating risk score...")
        progress_bar.progress(25)
        
        # Step 2: Search guidelines
        status_text.text("📚 Searching clinical guidelines...")
        progress_bar.progress(50)
        
        # Step 3: Search research
        status_text.text("🔬 Retrieving research evidence...")
        progress_bar.progress(75)
        
        # Step 4: Synthesize
        status_text.text("🧠 Synthesizing recommendations...")
        progress_bar.progress(90)
        
        # Call Agent Builder API
        try:
            headers = {
                "Authorization": f"ApiKey {ELASTIC_API_KEY}",
                "Content-Type": "application/json",
                "kbn-xsrf": "true"
            }
            
            # Construct natural language prompt for agent
            prompt = f"""
Assess cardiovascular disease risk for:
- {age_years}-year-old {"female" if gender == 1 else "male"}
- Height: {height}cm, Weight: {weight}kg (BMI: {bmi:.1f})
- Blood Pressure: {ap_hi}/{ap_lo} mmHg
- Cholesterol: {["normal", "above normal", "well above normal"][cholesterol-1]}
- Glucose: {["normal", "above normal", "well above normal"][gluc-1]}
- Smoker: {"Yes" if smoke else "No"}
- Alcohol: {"Yes" if alco else "No"}
- Physically Active: {"Yes" if active else "No"}

Please provide complete risk assessment with evidence-based recommendations.
"""
            
            
            base_url = "https://my-elasticsearch-project-f43eb6.kb.asia-southeast1.gcp.elastic.cloud"
            payload = {
    "agent_id": "cvd_7",
    "input": prompt
}

            response = requests.post(
                f"{base_url}/api/agent_builder/converse",
                headers=headers,
                json=payload,
                timeout=4000
            )
            progress_bar.progress(100)
            status_text.text("✅ Assessment complete!")
            st.write(response)
            
            if response.status_code == 200:
                result = response.json()
                st.session_state.risk_report_data = result
                # Parse agent response
                # (Structure depends on your agent's output format)
                agent_response = result.get('response', '')
                risk_data = result.get('risk_assessment', {})
                
                # Display results
                st.success("Risk Assessment Complete")
                
                # Risk Score Gauge
                st.subheader("Risk Score")
                
                risk_score = result['steps'][1]['results'][0]['data']['execution']['output']['data']['risk_score']
                print(f"Extracted Risk Score: {risk_score}")
                risk_percentage = int(risk_score * 100)
                risk_category = risk_data.get('risk_category', 'MODERATE')
                
                # Create gauge chart
                fig = go.Figure(go.Indicator(
                    mode="gauge+number+delta",
                    value=risk_percentage,
                    domain={'x': [0, 1], 'y': [0, 1]},
                    title={'text': "CVD Risk Score", 'font': {'size': 24}},
                    delta={'reference': 50},
                    gauge={
                        'axis': {'range': [None, 100], 'tickwidth': 1},
                        'bar': {'color': "darkblue"},
                        'bgcolor': "white",
                        'borderwidth': 2,
                        'bordercolor': "gray",
                        'steps': [
                            {'range': [0, 40], 'color': '#4caf50'},
                            {'range': [40, 70], 'color': '#ff9800'},
                            {'range': [70, 100], 'color': '#f44336'}
                        ],
                        'threshold': {
                            'line': {'color': "red", 'width': 4},
                            'thickness': 0.75,
                            'value': 90
                        }
                    }
                ))
                
                fig.update_layout(height=300)
                st.plotly_chart(fig, width='stretch')
                
                # Risk category box
                if risk_category == "HIGH":
                    st.markdown(f'<div class="risk-high"><b>⚠️ HIGH RISK</b> - {risk_percentage}% CVD Risk</div>', unsafe_allow_html=True)
                elif risk_category == "MODERATE":
                    st.markdown(f'<div class="risk-moderate"><b>⚡ MODERATE RISK</b> - {risk_percentage}% CVD Risk</div>', unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="risk-low"><b>✅ LOW RISK</b> - {risk_percentage}% CVD Risk</div>', unsafe_allow_html=True)
                
                # Top Risk Factors
                st.subheader("Primary Risk Factors")
                
                top_factors = risk_data.get('top_risk_factors', [])
                for i, factor in enumerate(top_factors[:5], 1):
                    with st.expander(f"{i}. {factor.get('feature_name', 'Unknown')} - {factor.get('interpretation', '')}"):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric("Value", factor.get('value', 'N/A'))
                            st.metric("Importance", f"{factor.get('importance', 0)*100:.1f}%")
                        with col2:
                            st.metric("SHAP Value", f"{factor.get('shap_value', 0):.3f}")
                            st.metric("Impact", factor.get('impact', 'unknown'))
                
                # Clinical Recommendations
                st.subheader("Evidence-Based Recommendations")
                st.markdown(agent_response['message'])
                
                # Create document in memory
                if st.session_state.risk_report_data:
                    
                    # Create the docx in-memory
                    doc = Document()
                    doc.add_heading('Clinical Risk Report', 0)
                    doc.add_paragraph(f"Score: {risk_percentage}")
                    doc.add_paragraph(agent_response['message'])
                    bio = io.BytesIO()
                    doc.save(bio)
                    
                    st.download_button(
                        label="Download Report",
                        data=bio.getvalue(),
                        file_name="risk_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
            else:
                st.error(f"Error: {response.status_code} - {response.text}")
                
        except Exception as e:
            st.error(f"Failed to get assessment: {str(e)}")
            st.info("Make sure your Agent Builder API is deployed and accessible")

else:
    # Show example/instructions when no assessment yet
    st.info("👈 Fill in patient information in the sidebar and click 'Assess Risk'")
    
    st.subheader("How It Works")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 1️⃣ ML Prediction")
        st.write("XGBoost model analyzes 14 patient features with SHAP interpretability")
    
    with col2:
        st.markdown("### 2️⃣ Evidence Retrieval")
        st.write("Agent searches clinical guidelines and research literature")
    
    with col3:
        st.markdown("### 3️⃣ Recommendations")
        st.write("Synthesizes evidence-based clinical recommendations with citations")
    
    st.divider()
    
    st.subheader("Sample Assessment")
    st.code("""
RISK SCORE: 78/100 (HIGH RISK)

PRIMARY RISK FACTORS:
1. Systolic BP 165 mmHg → 35% contribution
2. Cholesterol well above normal → 28% contribution  
3. Active smoking → 15% contribution

RECOMMENDATIONS:
1. BLOOD PRESSURE MANAGEMENT
   - Target: <130/80 mmHg
   - First-line: ACE inhibitor or ARB
   - Evidence: [2017 ACC/AHA HTN Guideline, Class I]
   
2. LIPID MANAGEMENT...
3. SMOKING CESSATION...
    """, language="text")